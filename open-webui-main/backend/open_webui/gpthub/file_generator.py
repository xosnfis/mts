"""
file_generator.py — Universal file generator for GPTHub.

Supports: .docx, .pdf (fpdf2), .txt, .xlsx, .csv
Files are stored in /app/backend/data/downloads.
"""

import re
import csv
import logging
import uuid
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)

DOWNLOAD_DIR = Path("/app/backend/data/downloads")

GENERATE_TAG_RE = re.compile(r"\[ACTION:GENERATE_FILE\(type=(\w+)\)\]", re.IGNORECASE)


def _ensure_dir() -> Path:
    DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)
    return DOWNLOAD_DIR


def strip_generate_tag(text: str) -> tuple[str, Optional[str]]:
    m = GENERATE_TAG_RE.search(text)
    if not m:
        return text, None
    fmt = m.group(1).lower()
    clean = GENERATE_TAG_RE.sub("", text).rstrip()
    return clean, fmt


def _md_to_lines(text: str) -> list[dict]:
    result = []
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            result.append({"type": "blank", "content": ""})
        elif line.startswith("### "):
            result.append({"type": "h3", "content": line[4:].strip()})
        elif line.startswith("## "):
            result.append({"type": "h2", "content": line[3:].strip()})
        elif line.startswith("# "):
            result.append({"type": "h1", "content": line[2:].strip()})
        elif re.match(r"^[-*]\s+", line):
            result.append({"type": "bullet", "content": re.sub(r"^[-*]\s+", "", line)})
        elif re.match(r"^\d+\.\s+", line):
            result.append({"type": "bullet", "content": re.sub(r"^\d+\.\s+", "", line)})
        elif re.match(r"^\*\*.+\*\*$", line.strip()):
            result.append({"type": "bold_line", "content": line.strip().strip("*")})
        else:
            result.append({"type": "normal", "content": line})
    return result


def _strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text


def _gen_txt(text: str, path: Path) -> None:
    path.write_text(text, encoding="utf-8")


def _gen_docx(text: str, path: Path) -> None:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for item in _md_to_lines(text):
        t = item["type"]
        c = item["content"]
        if t == "blank":
            doc.add_paragraph("")
        elif t == "h1":
            doc.add_heading(_strip_inline_md(c), level=1)
        elif t == "h2":
            doc.add_heading(_strip_inline_md(c), level=2)
        elif t == "h3":
            doc.add_heading(_strip_inline_md(c), level=3)
        elif t == "bullet":
            doc.add_paragraph(_strip_inline_md(c), style="List Bullet")
        elif t == "bold_line":
            p = doc.add_paragraph()
            p.add_run(c).bold = True
        else:
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*[^*]+\*\*)", c):
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(_strip_inline_md(part)).bold = True
                else:
                    p.add_run(_strip_inline_md(part))
    doc.save(str(path))


def _gen_pdf(text: str, path: Path) -> None:
    from fpdf import FPDF
    import os

    class PDF(FPDF):
        def footer(self):
            self.set_y(-12)
            self.set_font("DejaVu", size=8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"{self.page_no()}", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    font_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    bold_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
    ]
    regular_path = next((p for p in font_paths if os.path.exists(p)), None)
    bold_path = next((p for p in bold_paths if os.path.exists(p)), None)

    if regular_path:
        pdf.add_font("DejaVu", style="", fname=regular_path)
        pdf.add_font("DejaVu", style="B", fname=bold_path or regular_path)
        font_name = "DejaVu"
    else:
        font_name = "Helvetica"

    def sf(size=11, bold=False):
        pdf.set_font(font_name, style="B" if bold else "", size=size)

    for item in _md_to_lines(text):
        t = item["type"]
        c = _strip_inline_md(item["content"])
        if t == "blank":
            pdf.ln(4)
        elif t == "h1":
            pdf.ln(4); sf(18, True); pdf.set_text_color(30, 30, 30); pdf.multi_cell(0, 10, c); pdf.ln(2)
        elif t == "h2":
            pdf.ln(3); sf(14, True); pdf.set_text_color(50, 50, 50); pdf.multi_cell(0, 8, c); pdf.ln(1)
        elif t == "h3":
            pdf.ln(2); sf(12, True); pdf.set_text_color(70, 70, 70); pdf.multi_cell(0, 7, c)
        elif t == "bullet":
            sf(11); pdf.set_text_color(0, 0, 0); pdf.multi_cell(0, 6, f"  \u2022  {c}")
        elif t == "bold_line":
            sf(11, True); pdf.set_text_color(0, 0, 0); pdf.multi_cell(0, 6, c)
        else:
            sf(11); pdf.set_text_color(0, 0, 0); pdf.multi_cell(0, 6, c)

    pdf.output(str(path))


def _gen_csv(text: str, path: Path) -> None:
    lines = [l for l in text.split("\n") if l.strip()]
    with open(str(path), "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        for line in lines:
            if "|" in line:
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                if all(set(c) <= set("-: ") for c in cells):
                    continue
                writer.writerow(cells)
            else:
                writer.writerow([_strip_inline_md(line)])


def _gen_xlsx(text: str, path: Path) -> None:
    import pandas as pd
    lines = [l for l in text.split("\n") if l.strip()]
    rows = []
    for line in lines:
        if "|" in line:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue
            rows.append([_strip_inline_md(c) for c in cells])
        else:
            rows.append([_strip_inline_md(line)])
    if rows:
        max_cols = max(len(r) for r in rows)
        rows = [r + [""] * (max_cols - len(r)) for r in rows]
        df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
    else:
        df = pd.DataFrame({"content": [text]})
    df.to_excel(str(path), index=False, engine="openpyxl")


_EXT_MAP = {
    "txt":   ("txt",  _gen_txt),
    "docx":  ("docx", _gen_docx),
    "word":  ("docx", _gen_docx),
    "pdf":   ("pdf",  _gen_pdf),
    "csv":   ("csv",  _gen_csv),
    "xlsx":  ("xlsx", _gen_xlsx),
    "excel": ("xlsx", _gen_xlsx),
}

_FMT_LABELS = {
    "txt": "TXT", "docx": "Word", "word": "Word",
    "pdf": "PDF", "csv": "CSV", "xlsx": "Excel", "excel": "Excel",
}


def generate_file(text: str, fmt: str, base_name: str = "document") -> Optional[tuple[Path, str]]:
    entry = _EXT_MAP.get(fmt.lower())
    if not entry:
        log.warning("Unsupported format: %s", fmt)
        return None
    ext, generator = entry
    out_dir = _ensure_dir()
    filename = f"{base_name}_{uuid.uuid4().hex[:8]}.{ext}"
    file_path = out_dir / filename
    try:
        generator(text, file_path)
        log.info("Generated %s (%d bytes)", filename, file_path.stat().st_size)
        return file_path, filename
    except Exception as e:
        log.error("File generation error (%s): %s", fmt, e)
        return None


def make_download_link(filename: str, owui_base_url: str) -> str:
    url = f"{owui_base_url}/downloads/{filename}"
    label = _FMT_LABELS.get(filename.rsplit(".", 1)[-1].lower(), filename.rsplit(".", 1)[-1].upper())
    return f"[⬇️ Скачать ваш файл ({label})]({url})"
